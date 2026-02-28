import io
import json
import os
from dataclasses import dataclass, asdict
from typing import Dict, List, Optional

import pdfplumber
import streamlit as st
from docx import Document
from langchain_core.prompts import PromptTemplate
import openai
from transformers import T5Tokenizer, T5ForConditionalGeneration
from huggingface_hub import InferenceClient


# =========================
# Configurações Gerais
# =========================

APP_TITLE = "Analisador de Processos Judiciais com IA"
OPENAI_MODEL = "gpt-4o"  # usado apenas se optar por OpenAI

# Modelo local (offline) para sumarização jurídica em português
LOCAL_MODEL_NAME = "stjiris/t5-portuguese-legal-summarization"
_local_tokenizer: Optional[T5Tokenizer] = None
_local_model: Optional[T5ForConditionalGeneration] = None

# Modelo hospedado na Hugging Face (API)
HF_MODEL_NAME = "HuggingFaceH4/zephyr-7b-beta"  # pode ser alterado se preferir outro
_hf_client: Optional[InferenceClient] = None


# =========================
# Modelos de Dados
# =========================

@dataclass
class AnaliseProcesso:
    arquivo: str
    resumo_executivo: str
    pontos_criticos: str
    erros_gargalos: str
    sugestoes: str
    texto_completo: str


# =========================
# Utilitários de PDF
# =========================

def extrair_texto_pdf(file) -> str:
    """
    Extrai o texto de um PDF enviado via Streamlit (arquivo in-memory).
    Usa pdfplumber para uma extração mais limpa.
    """
    texto_total = []
    # `file` é um UploadedFile; precisamos de um buffer novo porque ele pode ser lido mais de uma vez
    with pdfplumber.open(io.BytesIO(file.getvalue())) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            texto_total.append(page_text)
    return "\n\n".join(texto_total).strip()


# =========================
# Prompt / IA (OpenAI + LangChain)
# =========================

ANALISE_TEMPLATE = """
Você é um advogado experiente no direito brasileiro. Receberá o inteiro teor de um processo judicial (petição inicial, contestação, decisões, etc.) em formato de texto plano.

Com base **exclusivamente** no texto fornecido, produza uma análise estruturada **em português** no seguinte formato JSON, SEM comentários adicionais, SEM texto fora do JSON, exatamente com estas chaves:

{{
  "resumo_executivo": "string",
  "pontos_criticos": "string",
  "erros_gargalos": "string",
  "sugestoes": "string"
}}

Instruções:
- **Resumo Executivo**: identifique, se possível, as partes do processo, o objeto da ação e o valor da causa. Seja claro e conciso.
- **Pontos Críticos**: liste os pedidos principais, os argumentos da parte contrária e pontos sensíveis que merecem atenção.
- **Erros/Gargalos**: aponte possíveis nulidades processuais, prazos perdidos, ausência de documentos essenciais, contradições ou fragilidades na argumentação.
- **Sugestões**: proponha teses jurídicas de defesa ou ataque, estratégias processuais e próximos passos recomendados.

IMPORTANTE:
- Se alguma informação não aparecer claramente no texto (por exemplo, valor da causa), explique essa limitação no respectivo campo, sem inventar dados.
- Respeite o formato JSON exatamente (aspas duplas, sem vírgulas sobrando, sem markdown).

Texto do processo:
------------------
{texto_processo}
"""

# Prompt para a barra de pesquisa / Vade Mecum inteligente
CONSULTA_LEI_TEMPLATE = """
Você é um advogado e professor de Direito Brasileiro (Penal, Civil, CLT e leis especiais).
Responda sempre em português do Brasil.

O usuário fará uma pergunta jurídica ou curiosa. Você deve:
- Identificar se é uma pergunta técnica ("tecnica") ou uma pergunta inusitada/meme ("meme").
- Trazer, se souber, o artigo exato de lei aplicável (código, lei especial, Constituição etc.).
- Caso NÃO consiga identificar com segurança um artigo específico, deixe claro que se trata de
  "interpretação baseada em princípios gerais" e não invente números de artigos.

Responda EXCLUSIVAMENTE no seguinte formato JSON (sem comentários, sem texto fora do JSON):
{{
  "tipo_pergunta": "tecnica ou meme",
  "lei_citada": "ex: Código Penal, art. 334, descaminho; ou 'interpretação baseada em princípios gerais'",
  "texto_artigo": "texto literal do dispositivo, se souber; senão deixe vazio ou explique a limitação",
  "comentario_ia": "comentário didático em português sobre o caso"
}}

Regras de tom:
- Se "tipo_pergunta" for "tecnica": use tom sério, objetivo e técnico.
- Se "tipo_pergunta" for "meme": use bom humor e leve sarcasmo, mas ainda citando leis reais que podem se aplicar
  (ex.: crimes contra a honra, difamação, injúria, etc.).

Pergunta do usuário:
--------------------
{pergunta_usuario}
"""

prompt_template = PromptTemplate(
    input_variables=["texto_processo"],
    template=ANALISE_TEMPLATE,
)


def get_openai_client_from_env() -> None:
    """
    Configura a chave da OpenAI a partir de env var ou secrets do Streamlit.
    Lança erro amigável se não encontrar.
    """
    # 1º tenta variável de ambiente
    api_key = os.getenv("OPENAI_API_KEY")

    # 2º tenta secrets do Streamlit, mas sem quebrar se não existir secrets.toml
    if not api_key:
        try:
            api_key = st.secrets["OPENAI_API_KEY"]
        except Exception:
            api_key = None
    if not api_key:
        st.error(
            "Chave da OpenAI não encontrada. "
            "Defina a variável de ambiente `OPENAI_API_KEY` ou use `.streamlit/secrets.toml`."
        )
        st.stop()
    openai.api_key = api_key


def get_local_model():
    """
    Carrega (lazy) o modelo T5 jurídico em português para sumarização local.
    """
    global _local_tokenizer, _local_model
    if _local_model is None or _local_tokenizer is None:
        with st.spinner("Carregando modelo local de sumarização jurídica (pode levar alguns minutos no primeiro uso)..."):
            _local_tokenizer = T5Tokenizer.from_pretrained(LOCAL_MODEL_NAME)
            _local_model = T5ForConditionalGeneration.from_pretrained(LOCAL_MODEL_NAME)
    return _local_tokenizer, _local_model


def get_hf_client() -> InferenceClient:
    """
    Configura o cliente da Hugging Face para usar a API de geração de texto.
    Procura o token em HUGGINGFACEHUB_API_TOKEN (env ou secrets).
    """
    global _hf_client
    if _hf_client is not None:
        return _hf_client

    token = os.getenv("HUGGINGFACEHUB_API_TOKEN")
    if not token:
        try:
            token = st.secrets["HUGGINGFACEHUB_API_TOKEN"]
        except Exception:
            token = None

    if not token:
        st.error(
            "Token da Hugging Face não encontrado. "
            "Defina a variável de ambiente `HUGGINGFACEHUB_API_TOKEN` "
            "ou use `.streamlit/secrets.toml` com essa chave."
        )
        st.stop()

    _hf_client = InferenceClient(model=HF_MODEL_NAME, token=token)
    return _hf_client


def chamar_ia_analise(texto_processo: str) -> Dict[str, str]:
    """
    Monta o prompt com LangChain (PromptTemplate) e chama a API da OpenAI,
    esperando um JSON com os campos especificados.
    """
    get_openai_client_from_env()

    prompt = prompt_template.format(texto_processo=texto_processo)

    try:
        response = openai.ChatCompletion.create(
            model=OPENAI_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": "Você é um advogado sênior, especialista em direito processual civil brasileiro.",
                },
                {
                    "role": "user",
                    "content": prompt,
                },
            ],
            temperature=0.1,
        )
        conteudo = response["choices"][0]["message"]["content"]
    except openai.error.RateLimitError as e:
        st.error(
            "A API da OpenAI retornou erro de limite de uso (RateLimitError). "
            "Você ultrapassou sua cota ou créditos. Acesse o painel da OpenAI, "
            "verifique o plano/faturamento ou use outra chave de API."
        )
        return {
            "resumo_executivo": "",
            "pontos_criticos": "",
            "erros_gargalos": f"Erro de cota na OpenAI: {e}",
            "sugestoes": "Regularize seu plano/cota da OpenAI ou configure uma nova chave antes de continuar a usar a análise automática.",
        }
    except openai.error.OpenAIError as e:
        st.error(f"Ocorreu um erro ao chamar a API da OpenAI: {e}")
        return {
            "resumo_executivo": "",
            "pontos_criticos": "",
            "erros_gargalos": f"Erro ao chamar a OpenAI: {e}",
            "sugestoes": "Tente novamente em alguns minutos. Se o problema persistir, revise a chave e as configurações da OpenAI.",
        }


def chamar_ia_local(texto_processo: str) -> Dict[str, str]:
    """
    Usa o modelo T5 jurídico em português (rodando localmente) para gerar
    uma análise estruturada sem depender de nenhuma API externa.
    Faz várias chamadas focadas (resumo, pontos críticos, erros, sugestões)
    para obter respostas mais didáticas.
    """
    tokenizer, model = get_local_model()

    # Limita o tamanho do texto de entrada para evitar estouro de memória
    max_chars = 8000
    texto_truncado = texto_processo[:max_chars]

    def _t5_responder(instrucao: str, texto: str, max_len: int = 256) -> str:
        prepared_text = "summarize: " + instrucao + "\n\n" + texto
        inputs = tokenizer.encode(
            prepared_text,
            return_tensors="pt",
            truncation=True,
            max_length=1024,
        )
        summary_ids = model.generate(
            inputs,
            num_beams=4,
            no_repeat_ngram_size=2,
            min_length=64,
            max_length=max_len,
            early_stopping=True,
        )
        return tokenizer.decode(summary_ids[0], skip_special_tokens=True).strip()

    resumo_executivo = _t5_responder(
        "Explique em linguagem jurídica clara quem são as partes, qual é o objeto da ação "
        "e, se constar no texto, o valor da causa. Seja didático e direto.",
        texto_truncado,
        max_len=220,
    )

    pontos_criticos = _t5_responder(
        "Liste de forma organizada os principais pedidos, os argumentos da parte autora "
        "e os argumentos da parte ré/contrária. Foque nos pontos realmente críticos do processo.",
        texto_truncado,
        max_len=260,
    )

    erros_gargalos = _t5_responder(
        "Identifique possíveis nulidades processuais, prazos eventualmente perdidos, "
        "ausência de documentos importantes e eventuais contradições no texto. "
        "Seja objetivo, em formato de parágrafo ou tópicos curtos.",
        texto_truncado,
        max_len=260,
    )

    sugestoes = _t5_responder(
        "Proponha teses jurídicas de defesa ou de ataque, estratégias processuais e "
        "próximos passos recomendados para o advogado, com linguagem prática.",
        texto_truncado,
        max_len=260,
    )

    return {
        "resumo_executivo": resumo_executivo,
        "pontos_criticos": pontos_criticos,
        "erros_gargalos": erros_gargalos,
        "sugestoes": sugestoes,
    }


def chamar_ia_hf(texto_processo: str) -> Dict[str, str]:
    """
    Usa um modelo hospedado na Hugging Face (Inference API) para gerar
    a mesma estrutura de análise que o GPT faria (JSON com 4 campos).
    """
    client = get_hf_client()

    # Limita o tamanho do texto enviado à API
    max_chars = 12000
    texto_truncado = texto_processo[:max_chars]

    prompt = ANALISE_TEMPLATE.format(texto_processo=texto_truncado)

    # Para modelos marcados como "conversational", usamos a API de chat_completion,
    # que é o equivalente ao ChatGPT na Hugging Face.
    try:
        resp = client.chat_completion(
            messages=[
                {
                    "role": "system",
                    "content": "Você é um advogado sênior, especialista em direito processual civil brasileiro.",
                },
                {
                    "role": "user",
                    "content": prompt,
                },
            ],
            max_tokens=800,
            temperature=0.15,
        )

        # Tenta extrair o conteúdo da primeira escolha, lidando com possíveis formatos
        conteudo = None
        try:
            # Pode ser objeto com atributos
            conteudo = resp.choices[0].message.content
        except Exception:
            try:
                # Ou dicionário
                conteudo = resp["choices"][0]["message"]["content"]
            except Exception:
                conteudo = str(resp)

    except Exception as e:
        st.error(f"Erro ao chamar a API da Hugging Face: {e}")
        return {
            "resumo_executivo": "",
            "pontos_criticos": "",
            "erros_gargalos": f"Erro ao chamar a Hugging Face: {e}",
            "sugestoes": "Verifique seu token da Hugging Face, o modelo configurado e tente novamente.",
        }

    # Tenta fazer o parse do JSON retornado
    try:
        dados = json.loads(conteudo)
    except json.JSONDecodeError:
        conteudo_limpo = conteudo.strip()
        inicio = conteudo_limpo.find("{")
        fim = conteudo_limpo.rfind("}")
        if inicio != -1 and fim != -1:
            try:
                dados = json.loads(conteudo_limpo[inicio : fim + 1])
            except Exception:
                dados = {
                    "resumo_executivo": conteudo_limpo,
                    "pontos_criticos": "",
                    "erros_gargalos": "",
                    "sugestoes": "",
                }
        else:
            dados = {
                "resumo_executivo": conteudo_limpo,
                "pontos_criticos": "",
                "erros_gargalos": "",
                "sugestoes": "",
            }

    return {
        "resumo_executivo": dados.get("resumo_executivo", "").strip(),
        "pontos_criticos": dados.get("pontos_criticos", "").strip(),
        "erros_gargalos": dados.get("erros_gargalos", "").strip(),
        "sugestoes": dados.get("sugestoes", "").strip(),
    }


def consultar_legislacao_openai(pergunta: str) -> Dict[str, str]:
    """
    Usa a API da OpenAI (GPT-4o) para responder dúvidas jurídicas pontuais
    na barra de pesquisa, retornando artigo de lei + comentário.
    """
    get_openai_client_from_env()

    prompt = CONSULTA_LEI_TEMPLATE.format(pergunta_usuario=pergunta)

    try:
        response = openai.ChatCompletion.create(
            model=OPENAI_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": "Você é um advogado sênior, especialista em Direito Brasileiro.",
                },
                {
                    "role": "user",
                    "content": prompt,
                },
            ],
            temperature=0.1,
        )
        conteudo = response["choices"][0]["message"]["content"]
    except openai.error.RateLimitError as e:
        st.error(
            "A API da OpenAI retornou erro de limite de uso (RateLimitError) ao consultar legislação. "
            "Você ultrapassou sua cota ou créditos. A consulta jurídica inteligente via OpenAI ficará limitada "
            "até que o plano/faturamento seja regularizado."
        )
        return {
            "tipo_pergunta": "tecnica",
            "lei_citada": "interpretação baseada em princípios gerais",
            "texto_artigo": "",
            "comentario_ia": f"Não foi possível chamar a OpenAI (RateLimitError: {e}). "
                             f"A consulta abaixo é apenas uma orientação genérica.",
        }
    except openai.error.OpenAIError as e:
        st.error(f"Ocorreu um erro ao chamar a API da OpenAI: {e}")
        return {
            "tipo_pergunta": "tecnica",
            "lei_citada": "interpretação baseada em princípios gerais",
            "texto_artigo": "",
            "comentario_ia": f"Erro ao consultar a legislação via OpenAI: {e}",
        }

    try:
        dados = json.loads(conteudo)
    except json.JSONDecodeError:
        conteudo_limpo = conteudo.strip()
        inicio = conteudo_limpo.find("{")
        fim = conteudo_limpo.rfind("}")
        if inicio != -1 and fim != -1:
            try:
                dados = json.loads(conteudo_limpo[inicio : fim + 1])
            except Exception:
                dados = {
                    "tipo_pergunta": "tecnica",
                    "lei_citada": "interpretação baseada em princípios gerais",
                    "texto_artigo": "",
                    "comentario_ia": conteudo_limpo,
                }
        else:
            dados = {
                "tipo_pergunta": "tecnica",
                "lei_citada": "interpretação baseada em princípios gerais",
                "texto_artigo": "",
                "comentario_ia": conteudo_limpo,
            }

    return {
        "tipo_pergunta": dados.get("tipo_pergunta", "tecnica").strip(),
        "lei_citada": dados.get("lei_citada", "").strip(),
        "texto_artigo": dados.get("texto_artigo", "").strip(),
        "comentario_ia": dados.get("comentario_ia", "").strip(),
    }


def consultar_legislacao_hf(pergunta: str) -> Dict[str, str]:
    """
    Usa a API da Hugging Face (modelo de chat) para responder dúvidas jurídicas pontuais.
    """
    client = get_hf_client()

    prompt = CONSULTA_LEI_TEMPLATE.format(pergunta_usuario=pergunta)

    try:
        resp = client.chat_completion(
            messages=[
                {
                    "role": "system",
                    "content": "Você é um advogado sênior, especialista em Direito Brasileiro.",
                },
                {
                    "role": "user",
                    "content": prompt,
                },
            ],
            max_tokens=800,
            temperature=0.15,
        )

        try:
            conteudo = resp.choices[0].message.content
        except Exception:
            try:
                conteudo = resp["choices"][0]["message"]["content"]
            except Exception:
                conteudo = str(resp)
    except Exception as e:
        st.error(f"Erro ao chamar a API da Hugging Face: {e}")
        return {
            "tipo_pergunta": "tecnica",
            "lei_citada": "interpretação baseada em princípios gerais",
            "texto_artigo": "",
            "comentario_ia": f"Erro ao consultar a legislação via Hugging Face: {e}",
        }

    try:
        dados = json.loads(conteudo)
    except json.JSONDecodeError:
        conteudo_limpo = conteudo.strip()
        inicio = conteudo_limpo.find("{")
        fim = conteudo_limpo.rfind("}")
        if inicio != -1 and fim != -1:
            try:
                dados = json.loads(conteudo_limpo[inicio : fim + 1])
            except Exception:
                dados = {
                    "tipo_pergunta": "tecnica",
                    "lei_citada": "interpretação baseada em princípios gerais",
                    "texto_artigo": "",
                    "comentario_ia": conteudo_limpo,
                }
        else:
            dados = {
                "tipo_pergunta": "tecnica",
                "lei_citada": "interpretação baseada em princípios gerais",
                "texto_artigo": "",
                "comentario_ia": conteudo_limpo,
            }

    return {
        "tipo_pergunta": dados.get("tipo_pergunta", "tecnica").strip(),
        "lei_citada": dados.get("lei_citada", "").strip(),
        "texto_artigo": dados.get("texto_artigo", "").strip(),
        "comentario_ia": dados.get("comentario_ia", "").strip(),
    }


def consultar_legislacao_local(pergunta: str) -> Dict[str, str]:
    """
    Consulta jurídica usando apenas o modelo local T5 (sem depender de APIs externas).
    Não tenta citar artigos exatos, apenas dá uma orientação geral.
    """
    q = pergunta.lower().strip()

    # Regra especial: "artigo 33" – uso popular costuma se referir ao art. 33 da Lei 11.343/2006 (tráfico de drogas)
    if "artigo 33" in q:
        texto_artigo = (
            "Lei nº 11.343/2006 (Lei de Drogas), art. 33, caput:\n"
            "\"Importar, exportar, remeter, preparar, produzir, fabricar, adquirir, vender, expor à venda, "
            "oferecer, ter em depósito, transportar, trazer consigo, guardar, prescrever, ministrar, entregar "
            "a consumo ou fornecer drogas, ainda que gratuitamente, sem autorização ou em desacordo com "
            "determinação legal ou regulamentar.\"\n\n"
            "Pena: reclusão de 5 a 15 anos e pagamento de 500 a 1.500 dias-multa."
        )
        comentario = (
            "No uso comum, quando alguém pergunta apenas \"artigo 33\" em contexto penal, geralmente está se "
            "referindo ao art. 33 da Lei de Drogas (Lei 11.343/2006), que tipifica o crime de tráfico ilícito "
            "de entorpecentes. Existem outros artigos 33 em outros diplomas (Código Penal, Código Civil, CLT etc.), "
            "mas este é o mais associado à expressão \"artigo 33\" em conversas jurídicas do dia a dia."
        )
        return {
            "tipo_pergunta": "tecnica",
            "lei_citada": "Lei 11.343/2006, art. 33 (tráfico ilícito de drogas)",
            "texto_artigo": texto_artigo,
            "comentario_ia": comentario,
        }

    # Demais casos: usa modelo local apenas para uma explicação geral, sem citar artigo exato
    tokenizer, model = get_local_model()

    instrucao = (
        "Você é um advogado brasileiro experiente. Explique de forma didática a dúvida abaixo, "
        "citando, se souber, quais ramos do direito se relacionam (Penal, Civil, Trabalho, etc.) "
        "e quais tipos de artigos geralmente tratam do tema, SEM inventar número de artigo específico. "
        "Se não tiver certeza, diga que se trata de interpretação baseada em princípios gerais."
    )

    prepared_text = instrucao + "\n\nPergunta do cliente:\n" + pergunta

    inputs = tokenizer.encode(
        prepared_text,
        return_tensors="pt",
        truncation=True,
        max_length=512,
    )

    summary_ids = model.generate(
        inputs,
        num_beams=4,
        no_repeat_ngram_size=2,
        min_length=96,
        max_length=256,
        early_stopping=True,
    )

    resposta = tokenizer.decode(summary_ids[0], skip_special_tokens=True).strip()

    return {
        "tipo_pergunta": "tecnica",
        "lei_citada": "interpretação baseada em princípios gerais (modelo local, sem artigo exato)",
        "texto_artigo": "",
        "comentario_ia": resposta,
    }


def montar_citacao_abnt(dados: Dict[str, str]) -> str:
    """
    Monta um texto simples para ser colado em petições,
    aproximando o estilo de citação ABNT (bloco recuado).
    """
    lei = dados.get("lei_citada", "").strip()
    artigo = dados.get("texto_artigo", "").strip()
    comentario = dados.get("comentario_ia", "").strip()

    corpo = artigo or comentario
    if not corpo:
        return ""

    linhas = corpo.splitlines()
    indent = "    "  # usuário pode aplicar recuo de 4 cm no Word
    corpo_indentado = "\n".join(indent + l for l in linhas if l.strip() != "")

    cabecalho = f"{lei}:\n" if lei else ""
    return f"{cabecalho}{corpo_indentado}"


# =========================
# Exportação para Word
# =========================

def gerar_docx(analise: AnaliseProcesso) -> bytes:
    """
    Gera um arquivo .docx em memória com base na análise.
    Retorna bytes prontos para download.
    """
    document = Document()

    document.add_heading("Relatório de Análise de Processo Judicial", level=1)
    document.add_paragraph(f"Arquivo analisado: {analise.arquivo}")

    document.add_heading("1. Resumo Executivo", level=2)
    document.add_paragraph(analise.resumo_executivo or "—")

    document.add_heading("2. Pontos Críticos", level=2)
    document.add_paragraph(analise.pontos_criticos or "—")

    document.add_heading("3. Erros / Gargalos", level=2)
    document.add_paragraph(analise.erros_gargalos or "—")

    document.add_heading("4. Sugestões e Teses", level=2)
    document.add_paragraph(analise.sugestoes or "—")

    document.add_heading("5. Texto Completo do Processo (extraído)", level=2)
    document.add_paragraph(analise.texto_completo or "—")

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# =========================
# UI / Streamlit
# =========================

def inicializar_estado():
    if "analises" not in st.session_state:
        st.session_state.analises: List[AnaliseProcesso] = []
    if "analise_selecionada" not in st.session_state:
        st.session_state.analise_selecionada: Optional[int] = None
    if "modo_analise" not in st.session_state:
        # Padrão: usar modelo local (sem API)
        st.session_state.modo_analise: str = "Modelo local (sem API)"
    if "ultima_consulta_lei" not in st.session_state:
        st.session_state.ultima_consulta_lei: Optional[Dict[str, str]] = None
    if "ultima_pergunta_lei" not in st.session_state:
        st.session_state.ultima_pergunta_lei: Optional[str] = None


def sidebar_layout():
    st.sidebar.title("📚 Processos & Histórico")
    st.sidebar.write("Faça upload de PDFs e acompanhe o histórico de análises.")

    # Escolha do modo de análise
    st.sidebar.subheader("Modo de análise")
    opcoes_modo = [
        "Modelo local (sem API)",
        "Hugging Face API (requere token)",
        "OpenAI (requer créditos)",
    ]
    try:
        idx_atual = opcoes_modo.index(st.session_state.modo_analise)
    except ValueError:
        idx_atual = 0

    modo = st.sidebar.radio(
        "Selecione como deseja analisar os PDFs:",
        options=opcoes_modo,
        index=idx_atual,
    )
    st.session_state.modo_analise = modo

    uploaded_files = st.sidebar.file_uploader(
        "Enviar PDFs do processo",
        type=["pdf"],
        accept_multiple_files=True,
    )

    if uploaded_files:
        if st.sidebar.button("Analisar PDFs com IA"):
            for f in uploaded_files:
                nome = f.name
                with st.spinner(f"Extraindo texto de `{nome}`..."):
                    texto = extrair_texto_pdf(f)

                if not texto:
                    st.warning(f"Nenhum texto extraído de `{nome}`.")
                    continue

                with st.spinner(f"Analisando `{nome}` com IA ({modo})..."):
                    if modo == "OpenAI (requer créditos)":
                        resultado = chamar_ia_analise(texto)
                    elif modo == "Hugging Face API (requere token)":
                        resultado = chamar_ia_hf(texto)
                    else:
                        resultado = chamar_ia_local(texto)

                analise = AnaliseProcesso(
                    arquivo=nome,
                    resumo_executivo=resultado["resumo_executivo"],
                    pontos_criticos=resultado["pontos_criticos"],
                    erros_gargalos=resultado["erros_gargalos"],
                    sugestoes=resultado["sugestoes"],
                    texto_completo=texto,
                )
                st.session_state.analises.append(analise)

            # Seleciona o último analisado por padrão
            if st.session_state.analises:
                st.session_state.analise_selecionada = len(st.session_state.analises) - 1
                st.success("Análise concluída para todos os PDFs enviados.")

    st.sidebar.markdown("---")
    st.sidebar.subheader("Histórico de análises")

    if st.session_state.analises:
        options = {
            f"{i+1}. {a.arquivo}": i for i, a in enumerate(st.session_state.analises)
        }
        label_default = (
            list(options.keys())[st.session_state.analise_selecionada]
            if st.session_state.analise_selecionada is not None
            else list(options.keys())[0]
        )

        escolha = st.sidebar.selectbox(
            "Selecione um processo analisado",
            options=list(options.keys()),
            index=list(options.keys()).index(label_default),
        )
        st.session_state.analise_selecionada = options[escolha]
    else:
        st.sidebar.info("Nenhuma análise realizada ainda.")

    # ------------------------------
    # Vade Mecum / Consulta Jurídica
    # ------------------------------
    st.sidebar.markdown("---")
    st.sidebar.subheader("📖 Vade Mecum Inteligente")

    pergunta_lei = st.sidebar.text_input(
        "Consultar legislação / dúvida jurídica",
        placeholder="Ex: Pena para descaminho; prazo de apelação cível; posso ser preso por fazer meme?",
    )

    if st.sidebar.button("Consultar legislação"):
        if not pergunta_lei.strip():
            st.sidebar.warning("Digite uma dúvida jurídica para consultar.")
        else:
            modo_atual = st.session_state.modo_analise
            with st.spinner(f"Consultando legislação ({modo_atual})..."):
                if modo_atual == "OpenAI (requer créditos)":
                    consulta = consultar_legislacao_openai(pergunta_lei)
                elif modo_atual == "Hugging Face API (requere token)":
                    consulta = consultar_legislacao_hf(pergunta_lei)
                else:
                    # Modo totalmente offline, usando apenas o modelo local
                    consulta = consultar_legislacao_local(pergunta_lei)

                st.session_state.ultima_consulta_lei = consulta
                st.session_state.ultima_pergunta_lei = pergunta_lei

    consulta_atual = st.session_state.ultima_consulta_lei
    if consulta_atual:
        st.sidebar.markdown("#### Resultado da consulta")

        titulo_artigo = consulta_atual.get("lei_citada") or "Referência legal"
        with st.sidebar.expander(titulo_artigo, expanded=True):
            texto_artigo = consulta_atual.get("texto_artigo") or "Interpretação baseada em princípios gerais."
            st.sidebar.info(texto_artigo)

        comentario = consulta_atual.get("comentario_ia") or ""
        if comentario:
            st.sidebar.markdown("💡 **Comentário da IA**")
            st.sidebar.write(comentario)

        citacao = montar_citacao_abnt(consulta_atual)
        if citacao:
            st.sidebar.markdown("##### Copiar para Petição")
            st.sidebar.text_area(
                "Texto formatado (selecione e copie):",
                value=citacao,
                height=180,
            )


def main_area():
    st.title(APP_TITLE)
    st.write(
        "Envie seus processos em PDF e receba um diagnóstico jurídico automático "
        "com resumo, pontos críticos, riscos e sugestões de teses."
    )

    if not st.session_state.analises:
        st.info("Comece fazendo upload de PDFs na barra lateral à esquerda.")
        return

    idx = st.session_state.analise_selecionada
    if idx is None or idx >= len(st.session_state.analises):
        st.warning("Selecione um processo na barra lateral para visualizar os insights.")
        return

    analise: AnaliseProcesso = st.session_state.analises[idx]

    st.subheader(f"Processo: {analise.arquivo}")

    aba_resumo, aba_pontos, aba_erros, aba_sugestoes, aba_texto = st.tabs(
        [
            "Resumo Executivo",
            "Pontos Críticos",
            "Erros / Gargalos",
            "Sugestões de Teses",
            "Texto Completo (extraído)",
        ]
    )

    with aba_resumo:
        st.markdown("### Resumo Executivo")
        st.write(analise.resumo_executivo or "—")

    with aba_pontos:
        st.markdown("### Pedidos e Pontos Críticos")
        st.write(analise.pontos_criticos or "—")

    with aba_erros:
        st.markdown("### Possíveis Nulidades / Gargalos")
        st.write(analise.erros_gargalos or "—")

    with aba_sugestoes:
        st.markdown("### Sugestões de Teses e Estratégias")
        st.write(analise.sugestoes or "—")

    with aba_texto:
        st.markdown("### Texto Completo Extraído do PDF")
        st.text_area(
            "Texto do processo (somente leitura)",
            value=analise.texto_completo,
            height=400,
        )

    st.markdown("---")
    st.subheader("Exportar relatório")

    docx_bytes = gerar_docx(analise)
    st.download_button(
        label="⬇️ Baixar relatório em Word (.docx)",
        data=docx_bytes,
        file_name=f"analise_{os.path.splitext(analise.arquivo)[0]}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def main():
    st.set_page_config(
        page_title=APP_TITLE,
        page_icon="⚖️",
        layout="wide",
    )

    inicializar_estado()
    sidebar_layout()
    main_area()


if __name__ == "__main__":
    main()

